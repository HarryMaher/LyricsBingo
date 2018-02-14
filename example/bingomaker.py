# -*- coding: UTF-8 -*-
#!/usr/bin/env python

'''
Spotify Playlist Bingo Card Maker!


Takes playlist.txt that contains a list of songs and artists*, then looks
up the lyrics from azlyrics. Scrapes the lyrics and outputs an out.txt
that has counts for how many songs have that word, and how many times
that word appears in all of the songs that have that word.


The most common 250 words (that appear in the highest number of songs) are 
put into a list of "common", from which we take a random sample of 25 of 
those words to put into a 5x5 grid to make the bingo sheet.

*Use this if you want to export a spotify playlist to text: 
http://spotlistr.herokuapp.com/#/export/spotify-playlist 


It's not very smart so different verb forms count as different words for now.
I'm sure there are libraries someone else wrote to help with that, but it's fine for now

'''
import sys
import requests
from bs4 import BeautifulSoup
import re
import time
from docx import Document
import random


#gets url from search
def get_href(url):
    headers = { 'User-Agent': 'Mozilla/5.0 (Windows NT 6.0; WOW64; rv:24.0) Gecko/20100101 Firefox/24.0' }
    response = requests.get(url, headers=headers)
    content = response.content
    soup = BeautifulSoup(content, "lxml")
    tag = [a['href'] for a in soup.find_all("a")]
    #returns 28th url which links to the (hopefully) correct lyrics, if it's to page 2, then give 29
    i = 28
    while  i < 40:
        if tag[i][:3]!='?q=':
            return tag[i]
        i+=1
    return tag[28] # gets tossed to not found & skipped cuz if it can't find a valid link after 40 give up...


def main(cardnum):

    with open("playlist.txt", "r") as f:
        data = f.readlines()

    searches= []
    forbidden_words = ["&amp;"]

    for lines in data:
        line = lines.split("-")[0].split(";")[0]
        line += lines.split("-")[1].replace("\n","")
        line = re.sub('\ \(.+?\)', '', line).replace(" ", "+").replace("&","")
        line = re.sub('\[.+?\]', '', line)
        searches.append(line)

    d = {}
    scrape_errors = []
    for search in searches:
        url= "https://search.azlyrics.com/search.php?q="+search
        print(url)
        url = get_href(url)
        time.sleep(3) # Sleep to not overwhelm the servers
        headers = { 'User-Agent': 'Mozilla/5.0 (Windows NT 6.0; WOW64; rv:24.0) Gecko/20100101 Firefox/24.0' } #what browser we're pretending to be
        try:
            response = requests.get(url, headers=headers)
            content = response.content
            soup = BeautifulSoup(content, "lxml")
            tag = soup.find("div", {"class": "col-xs-12 col-lg-8 text-center"})
            tag =str(tag)
            tag = tag.split('</div>')[6]
            tag = tag.replace("feat","")
            tag = re.sub('\[.+?\]', '', tag)
            tag = re.sub('\,', '', tag)
            tag = re.sub('\.', '', tag)
            tag = re.sub('\(', '', tag)
            tag = re.sub('\)', '', tag)
            tag = re.sub('\?', '', tag)
            tag = re.sub('\"', '', tag)
            lyrics = re.sub('<[^<]+?>', '', tag)
            lyrics = lyrics.split()
            thissong = []
            for word in lyrics:
                word = word.lower()
                if len(word)>3 and word not in forbidden_words:
                    if word not in thissong:
                        if word in d:
                            d[word][0] += 1
                            d[word][1] += 1
                        else:
                            d[word] = [1,1]
                        thissong.append(word)
                    else:
                        d[word][1]+=1
        except requests.exceptions.MissingSchema:
            print("NOT FOUND:", search)
            scrape_errors.append("NOT FOUND: "+search)

    #Write words to an empty out.txt for analysis
    with open('out.txt', 'w+') as fh:
        for err in scrape_errors:
            fh.write(err+"\n")
        fh.write("--------\n")
        for w in sorted(d, key=d.get, reverse=True):
            fh.write("{}, {}, {}\n".format(w,str(d[w][0]),str(d[w][1])))
            # would be more efficient to also build wordlist here. Meh.
    
    with open('out.txt', 'r') as f:
        wordlist = [line.split(',')[0] for line in f][len(scrape_errors)+1:251+len(scrape_errors)] # Top 250 words

    bingos = Document()
    
    for i in range(cardnum):
        table = bingos.add_table(rows=6, cols=5)
        table.style = 'TableGrid'
        heading_cells = table.rows[0].cells
        heading_cells[0].text = '\nB\n'
        heading_cells[1].text = '\nI\n'
        heading_cells[2].text = '\nN\n'
        heading_cells[3].text = '\nG\n'
        heading_cells[4].text = '\nO\n'

        words = random.sample(wordlist,25)  # get a random sample of 25 words from the wordlist 
        for r in range(1,6):
            for c in range(5):
                table.rows[r].cells[c].text="\n"+ words[r*5+c-5]+"\n" # Populate the chart with the words 
        
        table.rows[3].cells[2].text="\nFree!\n" # Middle is a Free space!
        p = bingos.add_paragraph('\n\n\n\n') # line breaks to separate sheets. Word 2016 gives two sheets a page
        bingos.save('bingo_sheet_output.docx')
        i+=1



if __name__ == "__main__":
    if len(sys.argv) > 2:
        cardnum = sys.argv[1]
    else:
        cardnum = 2
    main(cardnum)